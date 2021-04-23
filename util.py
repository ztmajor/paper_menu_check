# encoding: utf-8
from os import remove
from os.path import splitext
import win32com.client
from win32com.client import Dispatch, constants
import docx
import fitz
import os
import pythoncom


def get_word_com():
    try:
        doc_com = Dispatch('kwps.Application')
    except pythoncom.com_error:
        doc_com = Dispatch('Word.Application')

    return doc_com


def doc2docx(doc_path, docx_path):
    doc = get_word_com()
    doc_file = doc.Documents.Open(doc_path)
    doc_file.SaveAs(docx_path, 12, False, "", True, "", False, False, False, False)  # 转换后的文件,12代表转换后为docx文件
    # doc.SaveAs(r"F:\\***\\***\\appendDoc\\***.docx", 12) # 或直接简写
    # 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
    doc_file.Close()
    doc.Quit()


def read_docx(docx_path):
    return docx.Document(docx_path)


def doc2pdf(doc_file, pdf_file):
    wdFormatPDF = 17
    # word = Dispatch('kwps.Application')
    word = get_word_com()
    doc = word.Documents.Open(doc_file)
    doc.SaveAs(pdf_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()


def find_page_number(pdf_file, keyword):
    page_list = []
    with fitz.open(pdf_file) as doc:
        for idx, page in enumerate(doc, start=1):
            if keyword in page.getText():
                # print(f"page {idx} include keyword {keyword}!!")
                page_list.append(idx)
    return page_list


def check_alignment(docx_file):
    check_msg = []
    roma_nums_uppercase = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
    document = docx.Document(docx_file)
    for section in document.sections:
        for parg in section.footer.paragraphs:
            if parg.text:
                if parg.text in roma_nums_uppercase:
                    if str(parg.alignment) != 'CENTER (1)':
                        # print('页码 {} 未居中'.format(parg.text))
                        check_msg.append('页码 {} 未居中'.format(parg.text))
    return check_msg


def check_footer_nums(pdf_file, start_page, end_page):
    roma_nums_uppercase = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
    check_msg = []
    with fitz.open(pdf_file) as doc:
        for idx, page in enumerate(doc, start=1):
                if start_page <= idx < end_page:
                    cur_footer = roma_nums_uppercase[idx-start_page]
                    if '图目录' in page.getText():
                        cur_sec = '图目录'
                    elif '表目录' in page.getText():
                        cur_sec = '表目录'
                    elif '目录' in page.getText():
                        cur_sec = '目录'
                    if cur_footer in page.getText():
                        # print("{} 页码：{}".format(cur_sec, cur_footer))
                        check_msg.append("{} 页码：{}".format(cur_sec, cur_footer))
                    else:
                        # print("{} 缺少页码：{} 或者页码格式非大写罗马字符".format(cur_sec, cur_footer))
                        check_msg.append("{} 缺少页码：{} 或者页码格式非大写罗马字符".format(cur_sec, cur_footer))
    return check_msg


def check_footer(pdf, docx):
    '''
    不检查是否存在目录、表目录、图目录，默认都存在
    :param pdf: pdf file path
    :param docx: docx file path
    :return: None
    '''
    stand_page_list = find_page_number(pdf, '目录')
    p_list = find_page_number(pdf, '第 1 章')
    footer_num_msg = check_footer_nums(pdf, min(stand_page_list), min(p_list))
    alignment_msg = check_alignment(docx)
    return footer_num_msg + alignment_msg


def collect_infos(pdf_document, catalog_prefix=""):
    catalog_info, body_text_info, name_catalog = [], [], {}
    caption_id_index = {}
    next_chapter, cur_page = 1, -1
    cur_chapter_text = f"第{next_chapter}章"

    def extraId(text):
        return ''.join([i for i in text if i.isdigit() or i == '.'])

    for page in pdf_document:
        lines_str = page.getText().split("\n")
        for i in range(len(lines_str)):
            line = lines_str[i].strip()
            # 正常内容分布：第一行页眉，第二行页尾，第三行正文
            if line.isdigit() and i < 5:
                cur_page = int(line)

            infos = [item.strip() for item in line.split(' ') if item.strip()]
            if cur_page > 0 and i > 2:
                strip_line = ''.join(infos)
                if cur_chapter_text in strip_line:
                    next_chapter += 1
                    cur_chapter_text = f"第{next_chapter}章"

            if line.startswith(catalog_prefix) is False:
                continue

            if infos[0].strip() == catalog_prefix:
                # 检查该行第一个字是否是 指定的目录前缀，检查省略号（不一定是这种省略号）？
                if '..' in infos[-2]:
                    cur_id, cur_name, page_num = infos[1], ''.join(infos[2:-2]), infos[-1]
                    strip_name = catalog_prefix+cur_id+cur_name
                    # print(infos, cur_name, strip_name)
                    cur_index = len(catalog_info)
                    catalog_info.append({
                        'id': extraId(cur_id),
                        'page_num': page_num,
                        'name': strip_name,
                        'ori_line': line
                    })
                    name_catalog[strip_name] = cur_index
                elif not (line.endswith("：") or line.endswith("。")):
                    cur_id, cur_name = infos[1], ''.join(infos[2:])
                    # print(infos, next_chapter-1, cur_page, cur_name, lines_str[i], len(lines_str[i]))
                    strip_name = catalog_prefix + cur_id + cur_name
                    text_info = {
                            'id': extraId(cur_id),
                            'page_num': cur_page,
                            'chapter': next_chapter-1,
                            'name': strip_name,
                            'ori_line': line
                        }
                    if cur_id in caption_id_index:
                        body_text_info[caption_id_index[cur_id]] = text_info
                    else:
                        cur_index = len(body_text_info)
                        body_text_info.append(text_info)
                        caption_id_index[cur_id] = cur_index

    return catalog_info, name_catalog, body_text_info


def check_catalog_info(catalog_info):
    # id递增错误
    wrong_id_logs = []
    # 找不到页码
    wrong_page_num = []
    cur_chapter, cur_index, cur_page = 0, 1, 0
    for info in catalog_info:
        ids = info['id'].split('.')
        if cur_chapter != int(ids[0]):
            cur_chapter = int(ids[0])
            cur_index = 0

        cur_index += 1
        # todo: 此处转int很可能抛异常
        if cur_index != int(ids[1]):
            wrong_id_logs.append(
                f"目录中id不按序递增：{info['ori_line']}")

        if info['page_num'].isdigit() is False:
            wrong_page_num.append(
                f"目录中未显示页码：{info['ori_line']}")

    return wrong_id_logs + wrong_page_num


def get_captions(path):
    # 图表的名称应该是caption
    captions = []
    doc = docx.Document(path)
    for i in doc.paragraphs:
        if i.style.name == 'Caption':
            captions.append(i.text)

    return captions


def check_catalog(pdf_file, catalog_prefix=""):
    # 收集不同的错误日志并返回
    check_logs = []
    with fitz.open(pdf_file) as pdf_document:
        catalog_info, name_catalog, body_text_info = collect_infos(pdf_document, catalog_prefix)
        if len(catalog_info) == 0:
            return check_logs + [f"{catalog_prefix}目录未找到！"]

        # 检查目录页，判断id升序和页码信息
        check_logs += check_catalog_info(catalog_info)

        # 记录所在章节错误信息
        wrong_chapter_logs = []
        # id和目录中的不一致
        wrong_id_logs = []
        # 目录中页码错误
        wrong_page_num_logs = []
        # 目录中没有
        not_in_catalog_logs = []

        # 遍历全文检查图表实际位置
        for info in body_text_info:
            # print(info)
            if int(info['id'].split('.')[0]) != info['chapter']:
                wrong_chapter_logs.append(
                    f"章节和id不对应：第{info['page_num']}页中的 {info['ori_line']} 所在章节为 {info['chapter']}.")

            if info['name'] in name_catalog:
                catalog_data = catalog_info[name_catalog[info['name']]]
                if catalog_data['id'] != info['id']:
                    wrong_id_logs.append(
                        f"id与目录中的不一致: 第{info['page_num']}页中的 {info['ori_line']}，目录中id为{catalog_data['id']}.")
                if catalog_data['page_num'] != str(info['page_num']):
                    wrong_page_num_logs.append(
                        f"目录中页码错误: 第{info['page_num']}页中的 {info['ori_line']}，目录中页码为{catalog_data['page_num']}.")
            else:
                not_in_catalog_logs.append(
                    f"目录中没有: 第{info['page_num']}页中的 {info['ori_line']}.")
                # print(info['ori_line'])

    check_logs += wrong_chapter_logs
    check_logs += wrong_id_logs
    check_logs += wrong_page_num_logs
    check_logs += not_in_catalog_logs
    return check_logs


def collect_infos(pdf_document, catalog_prefix=""):
    catalog_info, body_text_info, name_catalog = [], [], {}
    caption_id_index = {}
    next_chapter, cur_page = 1, -1
    cur_chapter_text = f"第{next_chapter}章"

    def extraId(text):
        return ''.join([i for i in text if i.isdigit() or i == '.'])

    for page in pdf_document:
        lines_str = page.getText().split("\n")
        for i in range(len(lines_str)):
            line = lines_str[i].strip()
            # 正常内容分布：第一行页眉，第二行页尾，第三行正文
            if line.isdigit() and i < 5:
                cur_page = int(line)

            infos = [item.strip() for item in line.split(' ') if item.strip()]
            if cur_page > 0 and i > 2:
                strip_line = ''.join(infos)
                if cur_chapter_text in strip_line:
                    next_chapter += 1
                    cur_chapter_text = f"第{next_chapter}章"

            if line.startswith(catalog_prefix) is False:
                continue

            if infos[0].strip() == catalog_prefix:
                # 检查该行第一个字是否是 指定的目录前缀，检查省略号（不一定是这种省略号）？
                if '..' in infos[-2]:
                    cur_id, cur_name, page_num = infos[1], ''.join(infos[2:-2]), infos[-1]
                    strip_name = catalog_prefix+cur_id+cur_name
                    # print(infos, cur_name, strip_name)
                    cur_index = len(catalog_info)
                    catalog_info.append({
                        'id': extraId(cur_id),
                        'page_num': page_num,
                        'name': strip_name,
                        'ori_line': line
                    })
                    name_catalog[strip_name] = cur_index
                elif not (line.endswith("：") or line.endswith("。")):
                    cur_id, cur_name = infos[1], ''.join(infos[2:])
                    # print(infos, next_chapter-1, cur_page, cur_name, lines_str[i], len(lines_str[i]))
                    strip_name = catalog_prefix + cur_id + cur_name
                    text_info = {
                            'id': extraId(cur_id),
                            'page_num': cur_page,
                            'chapter': next_chapter-1,
                            'name': strip_name,
                            'ori_line': line
                        }
                    if cur_id in caption_id_index:
                        body_text_info[caption_id_index[cur_id]] = text_info
                    else:
                        cur_index = len(body_text_info)
                        body_text_info.append(text_info)
                        caption_id_index[cur_id] = cur_index

    return catalog_info, name_catalog, body_text_info


def check_catalog_info(catalog_info):
    # id递增错误
    wrong_id_logs = []
    # 找不到页码
    wrong_page_num = []
    cur_chapter, cur_index, cur_page = 0, 1, 0
    for info in catalog_info:
        ids = info['id'].split('.')
        if cur_chapter != int(ids[0]):
            cur_chapter = int(ids[0])
            cur_index = 0

        cur_index += 1
        # todo: 此处转int很可能抛异常
        if cur_index != int(ids[1]):
            wrong_id_logs.append(
                f"目录中id不按序递增：{info['ori_line']}")

        if info['page_num'].isdigit() is False:
            wrong_page_num.append(
                f"目录中未显示页码：{info['ori_line']}")

    return wrong_id_logs + wrong_page_num


def get_captions(path):
    # 图表的名称应该是caption
    captions = []
    doc = docx.Document(path)
    for i in doc.paragraphs:
        if i.style.name == 'Caption':
            captions.append(i.text)

    return captions


def check_catalog(pdf_file, catalog_prefix=""):
    # 收集不同的错误日志并返回
    check_logs = []
    with fitz.open(pdf_file) as pdf_document:
        catalog_info, name_catalog, body_text_info = collect_infos(pdf_document, catalog_prefix)
        if len(catalog_info) == 0:
            return check_logs + [f"{catalog_prefix}目录未找到！"]

        # 检查目录页，判断id升序和页码信息
        check_logs += check_catalog_info(catalog_info)

        # 记录所在章节错误信息
        wrong_chapter_logs = []
        # id和目录中的不一致
        wrong_id_logs = []
        # 目录中页码错误
        wrong_page_num_logs = []
        # 目录中没有
        not_in_catalog_logs = []

        # 遍历全文检查图表实际位置
        for info in body_text_info:
            # print(info)
            if int(info['id'].split('.')[0]) != info['chapter']:
                wrong_chapter_logs.append(
                    f"章节和id不对应：第{info['page_num']}页中的 {info['ori_line']} 所在章节为 {info['chapter']}.")

            if info['name'] in name_catalog:
                catalog_data = catalog_info[name_catalog[info['name']]]
                if catalog_data['id'] != info['id']:
                    wrong_id_logs.append(
                        f"id与目录中的不一致: 第{info['page_num']}页中的 {info['ori_line']}，目录中id为{catalog_data['id']}.")
                if catalog_data['page_num'] != str(info['page_num']):
                    wrong_page_num_logs.append(
                        f"目录中页码错误: 第{info['page_num']}页中的 {info['ori_line']}，目录中页码为{catalog_data['page_num']}.")
            else:
                not_in_catalog_logs.append(
                    f"目录中没有: 第{info['page_num']}页中的 {info['ori_line']}.")
                # print(info['ori_line'])

    check_logs += wrong_chapter_logs
    check_logs += wrong_id_logs
    check_logs += wrong_page_num_logs
    check_logs += not_in_catalog_logs
    return check_logs


if __name__ == '__main__':
    doc_path = os.getcwd() + '/paper/硕士学位论文正文_1.doc'
    pdf_path = doc_path.replace('doc', 'pdf')
    docx_path = doc_path.replace('doc', 'docx')
    # doc2docx(doc_path, docx_path)
    # for p in docx_file.paragraphs:
    #     print(p.text)
    # doc2pdf(doc_path, pdf_path)
    # stand_page_list = find_page_number(pdf_path, '目录')
    # p_list = find_page_number(pdf_path, '第 1 章')
    # print('final page:', max(p_list) - min(stand_page_list) + 1)
    msg = []
    msg = check_footer(pdf_path, docx_path)
    print(msg)

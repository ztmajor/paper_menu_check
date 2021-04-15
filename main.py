# encoding: utf-8
"""
@author: duke mei
@contact: mylovezju@163.com
@software: PyCharm
@file: main.py
@time: 2021/4/14 18:44
"""
import time
import os
import win32com
import win32com.client
import docx
from docx import Document
import argparse

# from util import doc2docx
cmd_opt = argparse.ArgumentParser(description='check')
cmd_opt.add_argument('-save_dir', default='.', help='result output root')


def find_footer(h):
    footers = [parg for section in h.sections for parg in section.footer.paragraphs if parg.text]
    footer_text = [footer.text for footer in footers]
    print(footer_text)


if __name__ == '__main__':
    file = docx.Document("paper/硕士学位论文正文_1.docx")
    # file = docx.Document('paper/test.docx')
    # print('段落:' + str(len(file.paragraphs)))
    # # for para in file.paragraphs:
    # #     print(para.text)
    # p = file.paragraphs
    # for i in range(len(file.paragraphs)):
    #     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
    find_footer(file)
